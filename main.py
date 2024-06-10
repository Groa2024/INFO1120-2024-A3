import sqlite3 as db
import pandas as pd


#Conectarse a la base de datos y obtener los registros de personas y rol (debe
#usar un inner join para obtener todos los datos).
def connect_db(db_file):
    conn = None
    try:
        conn = db.connect(db_file)
        print("Conexión a la base de datos SQLite exitosa")
        return conn
    except db.Error as e:
        print(f"Error al conectar a la base de datos SQLite: {e}")
    return conn

def fetch_persons_with_roles(conn):
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT personas.*, Salarios.rol AS rol 
            FROM personas 
            INNER JOIN Salarios ON personas.id_rol = id_rol
        """)
        rows = cursor.fetchall()
        return rows
    except db.Error as e:
        print(f"Error al obtener registros de la base de datos: {e}")
        return []

# Conectar a la base de datos
conn_db = connect_db('Sql_Data/db_personas.db')

# Obtener los registros de personas con sus roles
personas_con_roles = fetch_persons_with_roles(conn_db)

# Imprimir los registros obtenidos
for persona in personas_con_roles:
    print(persona)


#Traspase los registros de la base de datos a un dataframe de pandas que
#debería verse como la figura 1 (Anexo). 
def connect_db(db_file) :
    conn = None
    try :
        conn = db.connect(db_file)
        print("Conexión a la base de datos SQLite exitosa")
        return conn
    except db.Error as e:
        print(f"Error al conectar a la base de datos SQLite: {e}")
    return conn

def fetch_persons_with_roles(conn):
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT personas.*, Salarios.rol AS rol 
            FROM personas 
            INNER JOIN Salarios ON personas.id_rol = id_rol
        """)
        rows = cursor.fetchall()
        return rows
    except db.Error as e:
        print(f"Error al obtener registros de la base de datos: {e}")
        return []
#Genere un filtro que permita obtener singularmente a cualquiera de las
#personas dentro de los registros (Usando python-docx y el código de
#referencia realice esta implementación data.py)

#Obtenida la data, genere un contrato Word usando la libreria python-docx
#para esta persona usando la plantilla de código word_gen.py, este formato
#debera verse como la figura 2

#A partir del punto anterior, genere una nueva función que permita un rango
#de valores de inicio y final para la creacion de multiples documentos Word
#como se ve en la figura 3.
from docx import Document

def generar_documentos(rango_inicio, rango_final):
    for i in range(rango_inicio, rango_final + 1):
        document = Document()
        document.add_heading(f'Documento {i}', level=1)
        document.add_paragraph('Este es un documento de ejemplo.')

        document.save(f'documento_{i}.docx')
        print(f'Documento {i} creado.')


generar_documentos(1, 5)  # Esto creara 5 documentos desde documento_1.docx hasta documento_5.docx
#Utilizando la data anterior y la libreria matplotlib genere los siguiente
#graficos: 

#Grafico promedio sueldo por profesión (figura 4).
import matplotlib.pyplot as plt
import numpy as np

# Datos de ejemplo: profesiones y sus sueldos
profesiones = ['Administrador en sistema de bases de datos', 'Analisis de datos', 'Arquitecto de software', 'Cientifico de datos', 'Desarrollador web' ]
sueldos = [60000, 80000, 50000, 45000, 55000]  # en clp

# Calcular el sueldo promedio y la mediana
sueldo_promedio = np.mean(sueldos)
sueldo_mediana = np.median(sueldos)

# Crear el gráfico de barras
plt.figure(figsize=(10, 6))
plt.bar(profesiones, sueldos, color='skyblue', label='Sueldos')
plt.axhline(sueldo_promedio, color='red', linestyle='--', label=f'Promedio: ${sueldo_promedio:.2f}')
plt.axhline(sueldo_mediana, color='green', linestyle='-.', label=f'Mediana: ${sueldo_mediana:.2f}')
plt.ylabel('Sueldo ($)')
plt.xlabel('Profesión')
plt.title('Sueldo Promedio y Mediana por Profesión')
plt.legend()

# Mostrar el gráfico
plt.tight_layout()
plt.show()
#Grafico de tipo “tarta” que muestre la distribucion de profesiones
#(figura 5).
import matplotlib.pyplot as plt

# Datos de ejemplo: profesiones y sus frecuencias
profesiones = ['Administrador en sistema de bases de datos', 'Analisis de datos', 'Arquitecto de software', 'Cientifico de datos', 'Desarrollador web']
frecuencias = [20, 15, 10, 25, 30]  # Ejemplo de frecuencias, puedes ajustar según tus datos

# Crear el gráfico de tipo "tarta"
plt.figure(figsize=(8, 8))
plt.pie(frecuencias, labels=profesiones, autopct='%1.1f%%', startangle=140)
plt.title('Distribución de Profesiones')

# Mostrar el gráfico
plt.axis('equal')  # Para asegurar que el gráfico sea un circulo
plt.show()
#Grafico de conteo de profesionales por nacionalidad (figura 6).
import matplotlib.pyplot as plt
# Nacionalidades y conteo de profesionales
nacionalidades = ['Boliviana', 'Peruana', 'Chilena', 'Colombia', 'Argentina']
conteo_profesionales = [50, 30, 25, 20, 15]  # Ejemplo de conteo, puedes ajustar según tus datos

# Crear el gráfico de barras
plt.figure(figsize=(10, 6))
plt.bar(nacionalidades, conteo_profesionales, color='skyblue')
plt.xlabel('Nacionalidad')
plt.ylabel('Cantidad de Profesionales')
plt.title('Conteo de Profesionales por Nacionalidad')

# Mostrar el gráfico
plt.xticks(rotation=45)  # Rotar las etiquetas del eje x para mayor legibilidad
plt.tight_layout()
plt.show()
